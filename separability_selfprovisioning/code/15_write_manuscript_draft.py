#!/usr/bin/env python3
"""Build a conservative Paper 1 manuscript draft in Markdown and DOCX.

The manuscript intentionally follows the editor-review addendum rather than
the more optimistic earlier results package.
"""

from __future__ import annotations

import csv
import math
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "manuscript"
OUT.mkdir(parents=True, exist_ok=True)

MD_OUT = OUT / "paper1_manuscript_draft_revised.md"
DOCX_OUT = OUT / "paper1_manuscript_draft_revised.docx"


def read_csv(path: str) -> list[dict[str, str]]:
    with (ROOT / path).open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def fnum(x: str | float | int | None, digits: int = 3) -> str:
    if x is None or x == "":
        return ""
    try:
        val = float(x)
    except (TypeError, ValueError):
        return str(x)
    if not math.isfinite(val):
        return ""
    if abs(val) < 0.001 and val != 0:
        return f"{val:.2e}"
    return f"{val:.{digits}f}"


def pval(x: str | float | int | None) -> str:
    if x is None or x == "":
        return ""
    try:
        val = float(x)
    except (TypeError, ValueError):
        return str(x)
    if val < 0.001:
        return "<0.001"
    return f"{val:.3f}"


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    out = [
        "| " + " | ".join(headers) + " |",
        "|" + "|".join(["---"] * len(headers)) + "|",
    ]
    for row in rows:
        clean = [str(x).replace("|", "\\|") for x in row]
        out.append("| " + " | ".join(clean) + " |")
    return "\n".join(out)


sample = {r["item"]: r["value"] for r in read_csv("outputs/tables/table1_sample_summary_revised.csv")}
cat_part = read_csv("outputs/tables/table1_category_participation_revised.csv")
baseline = read_csv("outputs/tables/table2_common_sample_baseline.csv")
coeffs = read_csv("outputs/tables/table3_baseline_coefficients_margins.csv")
twopart = read_csv("outputs/tables/table5_two_part_model.csv")
addblock = read_csv("outputs/tables/tableE_add_one_block_diagnostics.csv")
village = read_csv("outputs/tables/tableF_village_fe_robustness.csv")
binary = read_csv("outputs/tables/tableG_binary_response_robustness.csv")
cats = read_csv("outputs/tables/tableI_category_variation_and_nsi_reframed.csv")
fixed_comp = read_csv("outputs/tables/tableJ_fixed_common_sample_robustness.csv")
bad = read_csv("outputs/tables/tableK_fixed_factors_bad_controls_robustness.csv")
price_diag = read_csv("outputs/tables/tableN_price_unit_value_diagnostics.csv")
iv = read_csv("outputs/tables/tableB_iv_diagnostics_appendix.csv")
market_int = read_csv("outputs/tables/tableA_market_friction_interactions_appendix.csv")


def by(rows: list[dict[str, str]], **kwargs: str) -> list[dict[str, str]]:
    out = []
    for r in rows:
        if all(r.get(k) == v for k, v in kwargs.items()):
            out.append(r)
    return out


def one(rows: list[dict[str, str]], **kwargs: str) -> dict[str, str]:
    hits = by(rows, **kwargs)
    if not hits:
        raise KeyError(kwargs)
    return hits[0]


cat_order = ["zhushi", "doulei", "roulei", "danlei", "nailei", "youzhi", "shucai", "shuiguo"]
cat_part_sorted = sorted(cat_part, key=lambda r: cat_order.index(r["food_category"]) if r["food_category"] in cat_order else 999)

table1_rows = [
    [
        r["food_category_label"],
        fnum(r["participation_rate"], 3),
        fnum(r["mean_self_suff_rate"], 3),
        fnum(r["mean_cons_monthly_jin"], 1),
        fnum(r["mean_selfprod_monthly_total"], 1),
    ]
    for r in cat_part_sorted
]

table2_rows = []
for outcome, label in [
    ("production_participation", "Participation"),
    ("log_selfprod_amount", "log(1 + self-produced quantity)"),
    ("ihs_selfprod_amount", "IHS self-produced quantity"),
]:
    for spec in ["M0", "M1", "M2", "M3"]:
        r = one(baseline, outcome=outcome, spec=spec)
        table2_rows.append([
            label,
            spec,
            r["n"],
            r["n_clusters"],
            fnum(r["hhcomp_wald_chisq"], 3),
            r["hhcomp_wald_df"],
            pval(r["hhcomp_wald_p"]),
        ])

add_part_rows = []
for label in [
    "B0_composition_category_year",
    "B1_plus_household_resources",
    "B1a_M1_plus_market",
    "B1b_M1_plus_GAEZ",
    "B1c_M1_plus_province_FE",
    "B2_full_market_GAEZ_province_FE",
    "B3_plus_unit_value_text",
]:
    r = one(addblock, outcome="production_participation", label=label)
    add_part_rows.append([
        label.replace("_", " "),
        fnum(r["wald_chisq"], 3),
        pval(r["wald_p"]),
    ])

table3_rows = []
for r in cats:
    table3_rows.append([
        r["food_category_label"],
        fnum(r["participation_rate"], 3),
        fnum(r["mean_self_suff_rate"], 3),
        fnum(r["nsi"], 3),
        pval(r["hhcomp_wald_p"]),
        pval(r["p_bh_fdr"]),
        r["main_text_status"],
    ])

rob_rows = []
for family in ["logit", "probit"]:
    r = one(binary, label="overall_M3", model_family=family)
    rob_rows.append([family.capitalize(), "Participation, M3", r["n"], r["n_clusters"], fnum(r["wald_chisq"], 3), pval(r["wald_p"])])
for outcome, label in [
    ("production_participation", "Village FE: participation"),
    ("log_selfprod_amount", "Village FE: log quantity"),
    ("ihs_selfprod_amount", "Village FE: IHS quantity"),
]:
    r = one(village, label="village_FE_M3_like", outcome=outcome)
    rob_rows.append([label, "Within-village comparison", r["n"], r["n_clusters"], fnum(r["wald_chisq"], 3), pval(r["wald_p"])])
for comp in ["proportion", "dependency", "counts"]:
    r = one(fixed_comp, composition_spec=comp, outcome="production_participation")
    rob_rows.append([f"Fixed sample: {comp}", "Participation", r["n"], r["n_clusters"], fnum(r["wald_chisq"], 3), pval(r["wald_p"])])
for label in ["fixed_factors_no_income_expense", "fixed_factors_no_income_expense_land_w99"]:
    r = one(bad, label=label, outcome="production_participation")
    rob_rows.append([label.replace("_", " "), "Participation", r["n"], r["n_clusters"], fnum(r["wald_chisq"], 3), pval(r["wald_p"])])

price_rows = [[r["diagnostic"], fnum(r["value"], 3), r["interpretation"]] for r in price_diag]

hh_coeff_rows = []
for term in ["household_size_reconstructed", "child_share", "elderly_share", "female_share"]:
    r = one(coeffs, outcome="production_participation", spec="M3", term=term)
    hh_coeff_rows.append([
        term,
        fnum(r["estimate"], 4),
        fnum(r["std_error_cluster"], 4),
        pval(r["p_value"]),
        r["direction"],
    ])

twopart_rows = []
for r in twopart:
    twopart_rows.append([
        r.get("part", r.get("model_part", "")),
        r["outcome"],
        r["n"],
        fnum(r.get("wald_chisq", r.get("hhcomp_wald_chisq", "")), 3),
        pval(r.get("wald_p", r.get("hhcomp_wald_p", ""))),
    ])

iv_rows = []
for r in iv:
    iv_rows.append([
        r.get("instrument", r.get("iv_name", r.get("iv_spec", ""))),
        fnum(r.get("min_first_stage_f", r.get("min_first_stage_F", r.get("min_f", ""))), 3),
        fnum(r.get("median_first_stage_f", r.get("median_first_stage_F", r.get("median_f", ""))), 3),
        r.get("weak_instrument_flag", r.get("weak_iv_flag", r.get("weak", ""))),
    ])

market_rows = []
for r in market_int:
    market_rows.append([
        r.get("friction_spec", r.get("model", "")),
        r["outcome"],
        fnum(r.get("interaction_wald_chisq", r.get("wald_chisq", "")), 3),
        pval(r.get("interaction_wald_p", r.get("wald_p", ""))),
    ])


manuscript = rf"""# Household Composition and Self-Provisioning: Multi-Category Evidence on Non-Separability in Rural China

**Draft date:** 2026-06-08  
**Status:** revised conservative manuscript draft based on the editor-review action plan  

## Abstract

Agricultural household models predict that, under complete markets and price-taking behavior, production decisions should be separable from household preferences and demographic composition. Most empirical tests of this prediction focus on labor demand or production input choices. This paper studies a different but substantively important margin: whether rural households enter self-provisioning for specific food categories. Using a pooled repeated cross-section of 3,565 rural Chinese households observed in 2023 or 2024, converted to 28,520 household-category observations across eight food categories, I test whether household size, child share, elderly share, and female share jointly predict category-specific self-provisioning. The preferred common-sample specification indicates that household composition predicts self-provisioning participation after controlling for household resources, local market access, agro-ecological suitability, purchase-side unit values, county text indicators, food-category fixed effects, province fixed effects, and survey-year fixed effects (Wald = 16.733, p = 0.002). The evidence is concentrated on the participation margin: full-sample transformed quantity outcomes are not significant in the preferred specification. However, the result is control-set sensitive. It is not significant in parsimonious specifications and does not survive a village fixed-effects participation-margin check, although log and IHS quantity margins become significant within villages. Category-level tests show strongest detectability for eggs, oils, vegetables, and fruits after false-discovery-rate adjustment, but these rankings should be interpreted as test-statistic detectability rather than economic magnitudes. The results provide cautious reduced-form evidence that household composition remains conditionally associated with rural food self-provisioning, while also underscoring the limits of cross-sectional separability tests without stronger panel or instrumental-variable identification.

**Keywords:** agricultural household model; separability; self-provisioning; household composition; rural China; food categories; market imperfections

## 1. Introduction

Rural households often make production and consumption decisions inside the same family enterprise. The classic separable agricultural household model offers a sharp benchmark: if markets are complete, all relevant prices are taken as given, and households can buy or sell goods and labor freely, production choices should be independent of household preferences and demographic composition. This recursive structure has made the agricultural household model one of the central tools in development and agricultural economics.

The empirical literature has most often tested this implication through farm labor demand. Benjamin (1992) tests whether household composition enters labor demand equations in Chinese agriculture. LaFave and Thomas (2016) revisit the same logic with longitudinal Indonesian data and reject the recursive model using within-household variation in composition and labor demand. These studies give the separability hypothesis a concrete empirical content: demographic variables that shape preferences and household labor endowments should not predict production choices once prices and fixed production factors are controlled.

This paper asks whether a related separability restriction holds for food self-provisioning. Instead of studying how much labor a household uses on the farm, I study whether the household produces part of its own food consumption in a category such as staples, eggs, vegetables, fruit, oils, beans, meat and aquatic products, or dairy. This margin matters because self-provisioning is common in rural settings and directly links production to consumption. It is also conceptually close to incomplete-market models: when buying and selling are frictionless at a common price, producing for own consumption is not economically distinct from producing for sale and buying food back. When transaction costs, quality differences, home-production preferences, or missing markets create a wedge between market purchase and household production, the decision to self-provision can become household-specific.

The empirical object is a pooled repeated cross-section of rural Chinese households. Each household appears in one survey year, 2023 or 2024, and contributes eight food-category rows. The final revised analysis file contains {sample.get("rows", "28,520")} rows, {sample.get("households", "3,565")} households, {sample.get("food_categories", "8")} food categories, {sample.get("villages_clusters", "361")} villages, {sample.get("counties", "44")} counties, and {sample.get("provinces", "9")} provinces. The design is not a household panel: each household identifier is observed in one year only. For this reason, the paper does not claim household fixed-effects identification. It implements a Benjamin-type reduced-form test in pooled cross-section and then evaluates how sensitive the result is to richer controls and village fixed effects.

The central finding is deliberately stated cautiously. Household composition significantly predicts self-provisioning participation in the preferred M3 common-sample specification, but not in simpler M0 or M1 specifications. The participation result appears after adding market-access, agro-ecological, and province controls, and remains significant when estimated with logit and probit. Yet it does not survive a village fixed-effects participation check. This pattern suggests that the data support a conditional association between household composition and self-provisioning, especially on the extensive margin, but do not justify a strong causal or structural claim that household demographics independently determine self-provisioning.

The paper makes three contributions. First, it extends separability testing from farm input demand to the food self-provisioning margin. Second, it shows that the association between household composition and self-provisioning is highly category-specific: eggs, oils, vegetables, and fruits remain significant after Benjamini-Hochberg false-discovery-rate correction, while dairy has too little variation to interpret substantively. Third, it provides a transparent account of weak evidence. Market-friction interactions do not provide strong mechanism evidence, candidate instrumental variables have weak first stages, purchase-side unit values are imperfect price proxies, and village fixed effects weaken the participation result.

## 2. Conceptual Framework

The separable agricultural household model starts from a household that derives utility from consumption goods and leisure while operating a production technology. When all output, input, labor, credit, and consumption markets are complete, and the household takes prices as given, production can be solved independently from consumption. The household first maximizes farm profit conditional on prices and fixed factors. Consumption and labor allocation are then chosen given income. In this recursive model, variables that shift preferences or household demographic needs should not enter production demand equations after conditioning on prices and fixed production factors.

The standard empirical restriction is therefore an exclusion restriction. Let \(y_{{hct}}\) denote a production-side outcome for household \(h\), category \(c\), and year \(t\). Let \(D_h\) be household composition and \(X_h\), \(M_v\), \(A_v\), and \(P_{{hct}}\) denote household resources, local market environment, agro-ecological conditions, and prices or price proxies. Under separability, conditional on the appropriate production-side controls, \(D_h\) should not predict \(y_{{hct}}\). The empirical test is whether the coefficients on household composition are jointly zero.

Self-provisioning requires one additional step in interpretation. The outcome in this paper is not total farm output or labor demand; it is whether households produce for their own consumption in a food category. In a world with no transaction costs and no quality or preference wedge between home-produced and purchased food, self-provisioning is not a distinct structural object: a household can sell output and buy the same food, or consume its own output, with no meaningful economic difference. In incomplete-market models with fixed or proportional transaction costs, however, households may face different effective buying and selling prices. The buy-sell price band can make non-participation or self-provisioning rational, and the household's demographic needs may become correlated with production-for-own-consumption decisions.

This logic implies a cautious interpretation. If household composition predicts self-provisioning, the recursive separability benchmark is strained. But the rejection can arise for several reasons: transaction costs and missing markets, home-produced food as a differentiated quality good, household-specific tastes for freshness or safety, or unobserved family orientation toward agriculture. The cross-sectional design cannot fully distinguish among these channels. I therefore treat market-friction interactions and instrumental-variable diagnostics as exploratory mechanism checks rather than the main identification basis.

## 3. Data and Variable Construction

### 3.1 Sample Structure

The analysis uses a pooled repeated cross-section of rural households. The sample contains one household-year per household identifier; no household appears in both survey years under the available identifier. The household-category file stacks eight food categories for each household. This yields {sample.get("rows", "28,520")} household-category observations from {sample.get("households", "3,565")} households.

The data cover nine provinces and 44 counties. Village identifiers are used for clustering and for a village fixed-effects robustness check. Because the sample is a repeated cross-section rather than a panel, the main specification uses food-category, year, and province fixed effects but not household fixed effects. Village fixed effects are feasible and are reported as a stringent robustness check that absorbs all time-invariant village-level differences in local markets, agricultural ecology, and food-production norms.

### 3.2 Outcomes

The primary outcome is an indicator for self-provisioning participation:

\[
Participation_{{hct}} = 1(SelfProducedQuantity_{{hct}} > 0).
\]

Two transformed quantity outcomes are used as secondary margins: \(\log(1 + SelfProducedQuantity_{{hct}})\) and \(asinh(SelfProducedQuantity_{{hct}})\). These transformations retain zero observations but are not scale-invariant; the paper therefore avoids elasticity-style interpretations and treats them as robustness outcomes. A self-sufficiency rate, defined as the share of category consumption supplied by self-production, is used in supplementary robustness checks.

### 3.3 Household Composition

The main household-composition variables are household size, child share, elderly share, and female share. They are intended to capture demographic structure rather than exogenous treatment. The sex-coding audit indicates that the gender code used to construct female share still requires manual codebook verification, so results involving female share should be interpreted with that caveat. The household roster is capped at eight members; 18 of 3,565 households reach this cap, or about 0.5 percent.

### 3.4 Controls

The preferred specification controls for household resources and fixed factors, including income and expenditure measures, agricultural and off-farm labor days, total sown area, household assets, and household-head characteristics. Because income and expenditure may be jointly determined with self-provisioning, I also report specifications that drop income and expenditure and condition only on fixed factors and price/unit-value controls.

Village and county context are measured using survey-based market-friction indices, lagged point-of-interest measures, agro-ecological suitability from GAEZ, and county-level text indicators. Food-category and survey-year fixed effects are included throughout the pooled models; province fixed effects are added in the richer specifications.

### 3.5 Prices and Unit Values

The available price proxy is a household purchase-side unit value, constructed from purchase expenditure divided by purchase quantity and measured in yuan per jin. It should not be interpreted as a pure exogenous market price or a farm-gate selling price. The distinction matters because the wedge between buying and selling prices is part of the incomplete-market mechanism that could generate self-provisioning. In the analysis file, 73.1 percent of unit values are observed from household purchase data and 26.9 percent are imputed using a hedonic model; the county-level hedonic model has an \(R^2\) of about 0.43 and log RMSE of about 0.72. Price robustness is therefore interpreted cautiously.

### 3.6 Descriptive Patterns

Table 1 summarizes category-level participation and self-sufficiency. Vegetables and staples have high participation rates, while dairy is almost degenerate. Dairy participation is only 0.13 percent, so dairy is excluded from substantive category interpretation. Vegetables and staples also require caution because participation is near the upper end of the distribution, leaving limited variation in the binary participation outcome.

**Table 1. Category-Level Self-Provisioning and Consumption**

{md_table(["Category", "Participation", "Self-sufficiency", "Mean consumption (jin/month)", "Mean self-produced (jin/month)"], table1_rows)}

## 4. Empirical Strategy

The main specification estimates

\[
y_{{hct}} = \alpha + D_h'\beta + X_h'\gamma + M_v'\delta + A_v'\theta + P_{{hct}}\pi + \mu_c + \lambda_t + \rho_p + \varepsilon_{{hct}},
\]

where \(y_{{hct}}\) is participation or a transformed quantity outcome, \(D_h\) is the vector of household-composition variables, \(X_h\) contains household resources and head characteristics, \(M_v\) contains market-access measures, \(A_v\) contains agro-ecological controls, \(P_{{hct}}\) is the unit-value proxy, \(\mu_c\) are food-category fixed effects, \(\lambda_t\) are year fixed effects, and \(\rho_p\) are province fixed effects. Standard errors are clustered at the village level.

The test of interest is a joint Wald test of

\[
H_0: \beta_{{size}} = \beta_{{child}} = \beta_{{elderly}} = \beta_{{female}} = 0.
\]

The models are organized as follows. M0 includes household composition, food-category fixed effects, and year fixed effects. M1 adds household resources and head controls. M2 adds market-friction controls, POI market measures, agro-ecological controls, and province fixed effects. M3 adds purchase-side unit values and county text controls.

The empirical design is a reduced-form separability test. It does not estimate a causal treatment effect of household composition. Household structure may be endogenous to migration, co-residence, agricultural orientation, and other unobserved factors. Under the separability null, however, demographic variables should be excluded from the production-side equation; observing a conditional association is therefore informative about the empirical adequacy of the separable benchmark, even if the mechanism behind rejection remains ambiguous.

## 5. Main Results

### 5.1 Baseline Wald Tests

Table 2 reports the M0-M3 sequence on the common M3 sample. The preferred M3 participation model rejects the joint exclusion of household composition (Wald = 16.733, p = 0.002). The same is not true in the parsimonious specifications: M0 has p = 0.178 and M1 has p = 0.106. The participation result becomes significant only after adding market, agro-ecological, and province controls.

The quantity outcomes show the opposite pattern. The log and IHS quantity outcomes are significant in M0 but collapse after household resources are added in M1. In M3, both full-sample quantity outcomes are insignificant. This contrast supports an extensive-margin interpretation in the preferred pooled specification, but it also reveals that the empirical pattern is sensitive to the control set.

**Table 2. Household-Composition Wald Tests Across Baseline Specifications**

{md_table(["Outcome", "Spec", "N", "Clusters", "Wald", "df", "p"], table2_rows)}

### 5.2 Which Controls Drive the Participation Result?

Table 3 decomposes the shift from M1 to M2. Adding market controls alone moves the participation test to p = 0.046; adding GAEZ controls alone gives p = 0.022; adding province fixed effects alone gives p = 0.012. The strongest intermediate block combines GAEZ and province fixed effects. This pattern means the M3 result should not be presented as invariant across reasonable specifications. Rather, it is a conditional association that emerges after accounting for regional, market-access, and agro-ecological heterogeneity.

**Table 3. Add-One-Block Diagnostics for Participation**

{md_table(["Specification block", "Wald", "p"], add_part_rows)}

### 5.3 Coefficient Patterns

In the preferred M3 participation model, household size is negatively associated with self-provisioning participation, elderly share is positively associated, child share is marginally positive, and female share is not statistically significant. The negative household-size coefficient may appear counterintuitive if size is treated only as labor availability; it may also reflect differences between larger non-agricultural households and smaller agriculturally oriented households after conditioning on resources. The elderly-share coefficient is more consistent with a household life-cycle or home-production interpretation, but it remains descriptive.

**Table 4. M3 Household-Composition Coefficients for Participation**

{md_table(["Variable", "Estimate", "Cluster SE", "p", "Direction"], hh_coeff_rows)}

## 6. Extensive and Intensive Margins

The two-part model separates entry into self-provisioning from conditional intensity among households that enter. Part 1 reproduces the participation result on all observations. Part 2 estimates log self-produced quantity only among positive self-provisioning observations. The conditional-intensity test is significant at the 5 percent level in this selected sample, but this result is descriptive because selection into the positive-production sample is itself a function of household composition.

**Table 5. Two-Part Model**

{md_table(["Part", "Outcome", "N", "Wald", "p"], twopart_rows)}

For this reason, the paper treats participation as the primary margin and conditional intensity as supplementary. The full-sample log and IHS models do not reject in M3, and transformed outcomes with many zeros are sensitive to scale and transformation choices.

## 7. Category Heterogeneity

Table 6 reports category-specific tests for participation. The Non-Separability Index (NSI) is defined as a category's household-composition Wald statistic divided by the mean Wald statistic across categories. This index is a relative detectability ranking, not an economic effect size. It combines coefficient magnitudes, residual variation, and precision. NSI values above one therefore mean only that the category's Wald statistic is above the category average.

After Benjamini-Hochberg FDR correction, eggs, oils, vegetables, and fruits remain significant at 5 percent. Beans is significant before correction but not after FDR correction. Oils require definition caution because the item-code audit is incomplete. Vegetables have the highest self-sufficiency rate but also high participation, so binary variation is compressed near the ceiling. Dairy is excluded from main category interpretation because almost no households self-provision dairy.

**Table 6. Category-Specific Detectability and Economic Importance**

{md_table(["Category", "Participation", "Self-sufficiency", "NSI", "Raw p", "BH FDR p", "Interpretation status"], table3_rows)}

The category results are substantively useful but should not be overread. Eggs have the strongest detectability ranking, but vegetables and staples are more economically important in self-sufficiency terms. This divergence illustrates why the paper reports participation rates and self-sufficiency alongside the Wald-based ranking.

## 8. Robustness and Sensitivity

Table 7 summarizes the main robustness checks. Logit and probit estimates of participation produce similar overall M3 Wald tests, indicating that the preferred participation result is not simply an artifact of the linear probability model. Fixed common-sample composition checks also support the participation result across proportion, dependency-ratio, and count measures.

The village fixed-effects check is more challenging for the main interpretation. With village fixed effects, the participation-margin Wald test is not significant (p = 0.171). The log and IHS quantity margins become significant under village fixed effects. This finding shifts the interpretation: the pooled M3 participation result is robust to functional form but not to fully absorbing village-level heterogeneity. The paper should therefore avoid language implying that the participation result is universally stable.

Dropping income and expenditure controls does not weaken the participation result. In fixed-factor specifications without income and expenditure, the participation Wald test remains significant. This reduces concern that the M3 participation result is produced only by conditioning on potentially endogenous financial variables.

**Table 7. Robustness Summary**

{md_table(["Check", "Outcome/specification", "N", "Clusters", "Wald", "p"], rob_rows)}

## 9. Mechanism Diagnostics

The preferred interpretation is that household composition is conditionally associated with self-provisioning participation. A stronger mechanism claim would require evidence that this association is amplified where markets are less complete or more costly to access. The available market-friction interactions do not provide that evidence. Interaction Wald tests using survey market friction, POI market friction, and combined market-friction measures are statistically weak across participation and quantity outcomes.

**Appendix Table A. Market-Friction Interaction Diagnostics**

{md_table(["Friction specification", "Outcome", "Wald", "p"], market_rows[:9])}

Candidate instrumental variables are also weak. Terrain-barrier and early-nighttime-light instruments have first-stage F-statistics far below conventional thresholds. IV results are therefore not used as evidence for causal identification.

**Appendix Table B. IV First-Stage Diagnostics**

{md_table(["Instrument", "Minimum F", "Median F", "Weak flag"], iv_rows)}

## 10. Price and Measurement Diagnostics

Table 8 summarizes the unit-value diagnostics. The observed measure is a purchase-side unit value, not a farm-gate output price. It may contain quality variation, quantity discounts, and selection into purchasing. Households that self-provision heavily may not purchase the category and therefore may have imputed rather than observed unit values. For these reasons, price controls are best understood as imperfect controls for local food-cost conditions rather than clean exogenous prices.

**Table 8. Price and Unit-Value Diagnostics**

{md_table(["Diagnostic", "Value", "Interpretation"], price_rows)}

Several data-definition limitations remain. The oils category requires item-code verification before strong substantive claims are made. Meat and aquatic products are aggregated with processed products in the current long file, limiting interpretation of protein-category heterogeneity. The current analysis-ready and cleaned long files no longer preserve item-level missingness for self-produced quantities, so a valid missing-exclusion versus missing-as-zero participation robustness check requires returning to raw item-level missing codes.

## 11. Discussion

The empirical evidence is best read as a disciplined but cautious separability test. The preferred pooled specification rejects the joint exclusion of household composition on the self-provisioning participation margin. The result is not driven by LPM functional form and remains when income and expenditure controls are removed. The category pattern also suggests that self-provisioning is not a uniform rural practice: it is more detectable for eggs, oils, vegetables, and fruits than for staples, meat/aquatic products, or dairy.

At the same time, the paper's limitations are central to its contribution. First, the data are pooled repeated cross-sections, not a panel. The design cannot absorb time-invariant household heterogeneity. Second, the participation result is control-set sensitive and does not survive village fixed effects. Third, market-friction interactions are weak, so the data do not strongly establish transaction costs as the specific channel. Fourth, unit values are imperfect price proxies. Fifth, category definitions for oils and meat/aquatic products require additional item-level validation.

These limitations imply that the paper should not claim a structural rejection of complete markets in the strong sense achieved by panel household fixed-effects designs. Instead, it contributes evidence on a less-studied margin: category-specific self-provisioning participation. The findings suggest that household demographic structure remains relevant for this margin even after rich controls, but that part of the relationship may reflect village-level heterogeneity, category measurement, and unobserved household orientation toward agriculture.

## 12. Conclusion

This paper studies whether household composition predicts food self-provisioning in rural China. Using a multi-category household-category dataset, I test the exclusion of household size, child share, elderly share, and female share from self-provisioning participation and quantity equations. In the preferred pooled specification, household composition significantly predicts participation, while full-sample quantity outcomes are weaker. The association is concentrated in selected categories and remains visible under logit and probit participation models, but it is sensitive to the inclusion of regional and village-level controls.

The main conclusion is therefore conservative: household composition conditionally predicts entry into self-provisioning, providing reduced-form evidence that the separable agricultural household benchmark is incomplete for this margin. The evidence does not establish a causal demographic effect or a clean market-friction mechanism. Future work should use panel data, validated item-level category definitions, better farm-gate and purchase price measures, and stronger instruments or natural experiments to distinguish market incompleteness from home-good quality preferences and unobserved household heterogeneity.

## References

Anderson, Michael L. 2008. "Multiple Inference and Gender Differences in the Effects of Early Intervention: A Reevaluation of the Abecedarian, Perry Preschool, and Early Training Projects." *Journal of the American Statistical Association* 103(484): 1481-1495.

Bellemare, Marc F., and Casey J. Wichman. 2020. "Elasticities and the Inverse Hyperbolic Sine Transformation." *Oxford Bulletin of Economics and Statistics* 82(1): 50-61.

Benjamin, Dwayne. 1992. "Household Composition, Labor Markets, and Labor Demand: Testing for Separation in Agricultural Household Models." *Econometrica* 60(2): 287-322.

Chen, Jiafeng, and Jonathan Roth. 2024. "Logs with Zeros? Some Problems and Solutions." Working paper.

de Janvry, Alain, Marcel Fafchamps, and Elisabeth Sadoulet. 1991. "Peasant Household Behaviour with Missing Markets: Some Paradoxes Explained." *Economic Journal* 101(409): 1400-1417.

Deaton, Angus. 1988. "Quality, Quantity, and Spatial Variation of Price." *American Economic Review* 78(3): 418-430.

Heckman, James J. 1979. "Sample Selection Bias as a Specification Error." *Econometrica* 47(1): 153-161.

Key, Nigel, Elisabeth Sadoulet, and Alain de Janvry. 2000. "Transactions Costs and Agricultural Household Supply Response." *American Journal of Agricultural Economics* 82(2): 245-259.

LaFave, Daniel, and Duncan Thomas. 2016. "Farms, Families, and Markets: New Evidence on Completeness of Markets in Agricultural Settings." *Econometrica* 84(5): 1917-1960.

Singh, Inderjit, Lyn Squire, and John Strauss, eds. 1986. *Agricultural Household Models: Extensions, Applications, and Policy*. Baltimore: Johns Hopkins University Press.
"""

MD_OUT.write_text(manuscript, encoding="utf-8")


def build_docx() -> None:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for margin in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(sec, margin, Inches(1))
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ]:
        st = styles[style_name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        st.font.size = Pt(size)
        st.font.bold = False
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.15

    def add_plain_title(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(26)
        run.font.bold = False
        run.font.color.rgb = RGBColor(0, 0, 0)

    def set_cell_text(cell, text: str, bold: bool = False) -> None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(str(text))
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(8.5)
        run.font.bold = bold

    def set_table_borders(table) -> None:
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = tblPr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tblPr.append(borders)
        for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            tag = "w:" + edge
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "DADCE0")

    def add_docx_table(headers: list[str], rows: list[list[str]]) -> None:
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table)
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            set_cell_text(hdr[i], h, bold=True)
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                set_cell_text(cells[i], val)
        doc.add_paragraph()

    add_plain_title("Household Composition and Self-Provisioning")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    r = subtitle.add_run("Multi-Category Evidence on Non-Separability in Rural China")
    r.font.name = "Arial"
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(85, 85, 85)
    meta = doc.add_paragraph("Revised conservative manuscript draft | 2026-06-08")
    meta.paragraph_format.space_after = Pt(12)

    current_table = None
    pending_table = None
    lines = manuscript.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("**Draft date:**") or line.startswith("**Status:**"):
            # Already represented in the clean metadata line below the title.
            pass
        elif line.startswith("**Keywords:**"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run("Keywords: ")
            run.bold = True
            rest = line.replace("**Keywords:**", "").strip()
            p.add_run(rest)
        elif line.startswith("**Table") or line.startswith("**Appendix Table"):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("**", ""))
            run.bold = True
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith("| "):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            headers = [x.strip() for x in table_lines[0].strip("|").split("|")]
            rows = []
            for tl in table_lines[2:]:
                rows.append([x.strip() for x in tl.strip("|").split("|")])
            add_docx_table(headers, rows)
            continue
        elif line.strip() == "":
            pass
        elif line.startswith("\\[") or line.startswith("$$") or line.startswith("\\]"):
            pass
        elif line.startswith("\\"):
            # Keep displayed equations readable as plain text in the DOCX draft.
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Inches(0.25)
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        i += 1

    doc.core_properties.title = "Household Composition and Self-Provisioning"
    doc.core_properties.subject = "Paper 1 revised manuscript draft"
    doc.core_properties.author = "Generated by Codex from local analysis outputs"
    doc.save(DOCX_OUT)


if __name__ == "__main__":
    build_docx()
    print(f"Wrote {MD_OUT}")
    print(f"Wrote {DOCX_OUT}")