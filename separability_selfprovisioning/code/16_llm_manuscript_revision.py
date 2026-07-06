#!/usr/bin/env python3
"""LLM-assisted manuscript revision workflow.

DeepSeek is used for token-heavy manuscript rewriting/integration.
Claude is used as an external reviewer.

API keys are read interactively or from environment variables and are never
written to disk.
"""

from __future__ import annotations

import getpass
import http.client
import json
import os
import re
import sys
import time
import urllib.error
import urllib.request
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "manuscript"
LOG_DIR = ROOT / "outputs" / "logs"
OUT.mkdir(parents=True, exist_ok=True)
LOG_DIR.mkdir(parents=True, exist_ok=True)

CURRENT_DRAFT = OUT / "paper1_manuscript_draft_revised.md"
EDITOR_PLAN = ROOT / "paper1_editor_review_and_action_plan.md"
ADDENDUM = ROOT / "outputs" / "reports" / "paper1_editor_revision_results_addendum.md"
RESULTS_PACKAGE = ROOT / "outputs" / "reports" / "paper1_revised_results_package.md"

DEEPSEEK_DRAFT = OUT / "paper1_manuscript_deepseek_rewrite.md"
CLAUDE_REVIEW = OUT / "paper1_claude_reviewer_comments.md"
FINAL_MD = OUT / "paper1_manuscript_llm_revised_final.md"
FINAL_DOCX = OUT / "paper1_manuscript_llm_revised_final.docx"
CALL_LOG = LOG_DIR / "llm_manuscript_revision_log.md"
DEEPSEEK_FRONT = OUT / "paper1_manuscript_deepseek_front.md"
DEEPSEEK_BACK = OUT / "paper1_manuscript_deepseek_back.md"
CLAUDE_FRONT = OUT / "paper1_claude_reviewer_front.md"
CLAUDE_BACK = OUT / "paper1_claude_reviewer_back.md"
CLAUDE_BACK_RESULTS = OUT / "paper1_claude_reviewer_back_results.md"
CLAUDE_BACK_DISCUSSION = OUT / "paper1_claude_reviewer_back_discussion.md"
FINAL_FRONT = OUT / "paper1_manuscript_llm_final_front.md"
FINAL_BACK = OUT / "paper1_manuscript_llm_final_back.md"


DEEPSEEK_URL = "https://api.deepseek.com/chat/completions"
DEEPSEEK_MODEL = "deepseek-v4-pro"
ANTHROPIC_URL = "https://api.anthropic.com/v1/messages"
ANTHROPIC_MODEL = "claude-sonnet-4-6"


def read_text(path: Path, max_chars: int | None = None) -> str:
    text = path.read_text(encoding="utf-8")
    if max_chars and len(text) > max_chars:
        return text[:max_chars] + "\n\n[TRUNCATED BY LOCAL WORKFLOW]\n"
    return text


def post_json(url: str, headers: dict[str, str], payload: dict, timeout: int = 240, retries: int = 2) -> dict:
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    last_error: Exception | None = None
    for attempt in range(retries + 1):
        req = urllib.request.Request(url, data=data, headers=headers, method="POST")
        try:
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                body = resp.read().decode("utf-8")
                return json.loads(body)
        except urllib.error.HTTPError as e:
            detail = e.read().decode("utf-8", errors="replace")
            raise RuntimeError(f"HTTP {e.code} from {url}: {detail[:2000]}") from e
        except (http.client.IncompleteRead, http.client.RemoteDisconnected, TimeoutError, urllib.error.URLError) as e:
            last_error = e
            if attempt < retries:
                time.sleep(4 + 3 * attempt)
                continue
            break
    raise RuntimeError(f"Request to {url} failed after retries: {last_error}") from last_error


def deepseek_chat(api_key: str, messages: list[dict[str, str]], max_tokens: int = 14000) -> str:
    return deepseek_chat_stream(api_key, messages, max_tokens=max_tokens)


def deepseek_chat_stream(api_key: str, messages: list[dict[str, str]], max_tokens: int = 8000) -> str:
    payload = {
        "model": DEEPSEEK_MODEL,
        "messages": messages,
        "temperature": 0.25,
        "max_tokens": max_tokens,
        "stream": True,
    }
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "Accept": "text/event-stream",
    }
    last_error: Exception | None = None
    for attempt in range(3):
        req = urllib.request.Request(DEEPSEEK_URL, data=data, headers=headers, method="POST")
        pieces: list[str] = []
        try:
            with urllib.request.urlopen(req, timeout=360) as resp:
                for raw in resp:
                    line = raw.decode("utf-8", errors="replace").strip()
                    if not line or not line.startswith("data:"):
                        continue
                    event = line[len("data:"):].strip()
                    if event == "[DONE]":
                        break
                    try:
                        obj = json.loads(event)
                    except json.JSONDecodeError:
                        continue
                    choice = obj.get("choices", [{}])[0]
                    delta = choice.get("delta", {})
                    text = delta.get("content") or choice.get("message", {}).get("content") or ""
                    if text:
                        pieces.append(text)
                text = "".join(pieces)
                if text.strip():
                    return text
                raise RuntimeError("DeepSeek stream returned no text.")
        except urllib.error.HTTPError as e:
            detail = e.read().decode("utf-8", errors="replace")
            raise RuntimeError(f"HTTP {e.code} from DeepSeek: {detail[:2000]}") from e
        except (http.client.IncompleteRead, TimeoutError, urllib.error.URLError, RuntimeError) as e:
            last_error = e
            if pieces:
                return "".join(pieces)
            time.sleep(4 + 3 * attempt)
    raise RuntimeError(f"DeepSeek streaming request failed after retries: {last_error}") from last_error


def get_or_generate(path: Path, label: str, fn) -> str:
    if path.exists() and path.stat().st_size > 100:
        print(f"Using cached {label}: {path}", flush=True)
        return path.read_text(encoding="utf-8")
    text = clean_markdown(fn())
    path.write_text(text, encoding="utf-8")
    print(f"Wrote {label}: {path}", flush=True)
    return text


def claude_review(api_key: str, prompt: str, max_tokens: int = 6000) -> str:
    payload = {
        "model": ANTHROPIC_MODEL,
        "max_tokens": max_tokens,
        "temperature": 0.2,
        "messages": [{"role": "user", "content": prompt}],
    }
    out = post_json(
        ANTHROPIC_URL,
        {
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
        payload,
        timeout=300,
    )
    try:
        return "\n".join(part.get("text", "") for part in out["content"] if part.get("type") == "text")
    except Exception as e:
        raise RuntimeError(f"Unexpected Claude response shape: {json.dumps(out)[:2000]}") from e


def clean_markdown(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^```(?:markdown)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    return text.strip() + "\n"


def build_deepseek_rewrite_prompt(current: str, editor: str, addendum: str, results: str) -> list[dict[str, str]]:
    system = (
        "You are a senior agricultural/development economics manuscript writer. "
        "Rewrite a complete conservative English paper draft using only the supplied evidence. "
        "Do not invent data, specifications, significance, citations, or causal claims. "
        "Keep all reviewer caveats explicit. Output Markdown only, with no code fences."
    )
    user = f"""
Task: Produce a polished full manuscript draft for Paper 1.

Mandatory stance:
- This is a pooled repeated-cross-section reduced-form separability test, not a causal paper.
- Main wording must be conservative: M3 participation is significant, but the result is control-set sensitive and does not survive the village-FE participation check.
- Quantity margins are secondary and sensitive.
- NSI is a relative Wald-statistic detectability ranking, not an economic magnitude.
- Dairy must be excluded from substantive category interpretation.
- Price is purchase-side unit value, not farm-gate price.
- Market-friction interactions and IV diagnostics are weak/exploratory.

Desired output:
- Complete English manuscript in Markdown.
- Include Abstract, Introduction, Conceptual Framework, Data, Empirical Strategy, Results, Robustness/Sensitivity, Mechanism Diagnostics, Discussion, Conclusion, References.
- Include concise Markdown tables when useful, but do not create new results.
- Use an academic but plain style suitable for AJAE/ERAE/Food Policy revision.

Current Codex draft:
{current}

Editor/reviewer action plan:
{editor}

Editor-revision empirical addendum:
{addendum}

Earlier revised results package:
{results}
"""
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def build_claude_review_prompt(draft: str, addendum: str) -> str:
    return f"""
You are a skeptical reviewer for AJAE, ERAE, Food Policy, or AEPP.

Review the manuscript below. Focus on whether the draft:
1. overclaims separability rejection or causality;
2. transparently handles control-set sensitivity and the village fixed-effects result;
3. treats NSI correctly as detectability rather than economic magnitude;
4. handles price/unit-value limitations correctly;
5. handles two-part model selection and weak IV/mechanism evidence correctly;
6. has missing tables, confusing structure, or unclear contribution;
7. needs wording changes before being circulated.

Return an actionable review memo with:
- Major issues;
- Minor issues;
- Specific passages or claims to revise;
- A concise recommended revision strategy.

Do not rewrite the whole manuscript.

Empirical addendum to use as ground truth:
{addendum}

Manuscript draft to review:
{draft}
"""


def build_claude_segment_review_prompt(segment_name: str, segment: str, addendum: str) -> str:
    return f"""
You are a skeptical reviewer for AJAE, ERAE, Food Policy, or AEPP.

Review this manuscript segment only: {segment_name}.

Focus on whether the segment:
1. overclaims separability rejection or causality;
2. transparently handles control-set sensitivity and the village fixed-effects result where relevant;
3. treats NSI as detectability rather than economic magnitude where relevant;
4. handles price/unit-value limitations where relevant;
5. correctly describes weak IV/mechanism evidence where relevant;
6. has unclear contribution, missing caveats, or wording that should be softened.

Return an actionable reviewer memo with major issues, minor issues, and specific revision instructions.

Use this empirical addendum as ground truth:
{addendum}

Segment to review:
{segment}
"""


def build_final_integration_prompt(deepseek_draft: str, claude_comments: str, addendum: str) -> list[dict[str, str]]:
    system = (
        "You are a senior agricultural economics coauthor revising a manuscript after reviewer comments. "
        "Integrate the reviewer feedback into a complete final Markdown manuscript. "
        "Do not invent results or remove necessary caveats. Output Markdown only, no code fences."
    )
    user = f"""
Revise the manuscript below in response to the Claude reviewer memo.

Ground rules:
- Keep the paper conservative and reduced-form.
- Preserve the key numerical results from the empirical addendum.
- Make the introduction and conclusion candid about the village-FE participation result.
- Clearly mark mechanism and IV evidence as weak/exploratory.
- Do not hide data-definition limitations.
- Keep the draft coherent and ready for human coauthor editing.

Empirical addendum:
{addendum}

Claude reviewer memo:
{claude_comments}

Draft to revise:
{deepseek_draft}
"""
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def build_segment_prompt(segment_name: str, current: str, editor: str, addendum: str, results: str) -> list[dict[str, str]]:
    system = (
        "You are a senior agricultural/development economics manuscript writer. "
        "Rewrite only the requested manuscript segment in polished academic English. "
        "Use only the supplied evidence, preserve caveats, and output Markdown only with no code fences."
    )
    if segment_name == "front":
        target = (
            "Write these sections only: Title, Abstract, Keywords, 1. Introduction, "
            "2. Conceptual Framework, 3. Data and Variable Construction, 4. Empirical Strategy."
        )
    elif segment_name == "back":
        target = (
            "Write these sections only: 5. Main Results, 6. Extensive and Intensive Margins, "
            "7. Category Heterogeneity, 8. Robustness and Sensitivity, 9. Mechanism Diagnostics, "
            "10. Price and Measurement Diagnostics, 11. Discussion, 12. Conclusion, References."
        )
    else:
        target = f"Write the requested segment only: {segment_name}."
    user = f"""
Task: {target}

Mandatory empirical stance:
- Pooled repeated cross-section, not panel or causal identification.
- M3 participation result is significant, but M0/M1 are not and village-FE participation is not significant.
- Quantity margins are secondary/sensitive.
- NSI is relative Wald-statistic detectability, not economic magnitude.
- Dairy is excluded from substantive category interpretation.
- Price is purchase-side unit value, not farm-gate price.
- Market-friction interactions and IV are weak/exploratory.

Use the current draft as structure and the addendum as ground truth. Keep prose concise but complete.

Current draft:
{current}

Editor action plan excerpt:
{editor}

Empirical addendum:
{addendum}

Earlier results package:
{results}
"""
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def build_segment_revision_prompt(segment_name: str, segment_text: str, full_draft: str, claude_comments: str, addendum: str) -> list[dict[str, str]]:
    system = (
        "You are a senior agricultural economics coauthor revising one manuscript segment after reviewer comments. "
        "Revise only the supplied segment. Output Markdown only, no code fences."
    )
    user = f"""
Revise this manuscript segment in response to the Claude reviewer memo.

Segment: {segment_name}

Ground rules:
- Preserve all necessary caveats.
- Keep the result numbers consistent with the empirical addendum.
- Do not introduce new empirical claims or new citations beyond the supplied text.
- If the reviewer asks for a change that cannot be supported by the data, state the limitation instead.

Empirical addendum:
{addendum}

Claude reviewer memo:
{claude_comments}

Full draft context:
{full_draft[:24000]}

Segment to revise:
{segment_text}
"""
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def add_docx_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "DADCE0")
        borders.append(el)
    table._tbl.tblPr.append(borders)

    def put(cell, text: str, bold: bool = False) -> None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(text))
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
        run.font.bold = bold

    for i, h in enumerate(headers):
        put(table.rows[0].cells[i], h, True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            put(cells[i], val)
    doc.add_paragraph()


def markdown_to_docx(markdown: str, out_path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ]:
        st = doc.styles[style_name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        st.font.size = Pt(size)
        st.font.bold = False
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.15

    lines = markdown.splitlines()
    i = 0
    title_done = False
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("| "):
            block = []
            while i < len(lines) and lines[i].startswith("|"):
                block.append(lines[i])
                i += 1
            if len(block) >= 2:
                headers = [x.strip() for x in block[0].strip("|").split("|")]
                rows = [[x.strip() for x in b.strip("|").split("|")] for b in block[2:]]
                add_docx_table(doc, headers, rows)
            continue
        if line.startswith("# "):
            if not title_done:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)
                r = p.add_run(line[2:].strip())
                r.font.name = "Arial"
                r.font.size = Pt(26)
                r.font.bold = False
                title_done = True
            else:
                doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif not line.strip():
            pass
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line).strip(), style="List Number")
        else:
            txt = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            txt = re.sub(r"\*(.*?)\*", r"\1", txt)
            p = doc.add_paragraph(txt)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        i += 1

    doc.core_properties.title = "Household Composition and Self-Provisioning"
    doc.core_properties.author = "Codex + DeepSeek rewrite + Claude review workflow"
    doc.save(out_path)


def main() -> int:
    current = read_text(CURRENT_DRAFT)
    editor = read_text(EDITOR_PLAN, max_chars=42000)
    addendum = read_text(ADDENDUM)
    results = read_text(RESULTS_PACKAGE)

    deepseek_key = os.environ.get("DEEPSEEK_API_KEY") or getpass.getpass("DeepSeek API key: ")
    anthropic_key = os.environ.get("ANTHROPIC_API_KEY") or getpass.getpass("Anthropic API key: ")
    if not deepseek_key or not anthropic_key:
        raise SystemExit("Both API keys are required.")

    log_lines = [
        "# LLM Manuscript Revision Log",
        "",
        f"Generated at: {time.strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "- DeepSeek role: token-heavy manuscript rewrite and final integration.",
        "- Claude role: skeptical reviewer.",
        "- API keys were read interactively/environment and not written to disk.",
        "",
    ]

    print("Calling DeepSeek for segmented front-half rewrite...", flush=True)
    front = get_or_generate(
        DEEPSEEK_FRONT,
        "DeepSeek front rewrite",
        lambda: deepseek_chat(deepseek_key, build_segment_prompt("front", current, editor, addendum, results), max_tokens=4200),
    )
    print("Calling DeepSeek for segmented back-half rewrite...", flush=True)
    back = get_or_generate(
        DEEPSEEK_BACK,
        "DeepSeek back rewrite",
        lambda: deepseek_chat(deepseek_key, build_segment_prompt("back", current, editor, addendum, results), max_tokens=5200),
    )
    ds_draft = clean_markdown(front + "\n\n" + back)
    DEEPSEEK_DRAFT.write_text(ds_draft, encoding="utf-8")
    log_lines.append(f"- DeepSeek rewrite written: `{DEEPSEEK_DRAFT.relative_to(ROOT)}` ({len(ds_draft)} chars).")

    print("Calling Claude as reviewer for front half...", flush=True)
    review_front = get_or_generate(
        CLAUDE_FRONT,
        "Claude reviewer memo front",
        lambda: claude_review(anthropic_key, build_claude_segment_review_prompt("front half", front, addendum), max_tokens=2600),
    )
    print("Calling Claude as reviewer for back half...", flush=True)
    if CLAUDE_BACK.exists() and CLAUDE_BACK.stat().st_size > 100:
        review_back = CLAUDE_BACK.read_text(encoding="utf-8")
        print(f"Using cached Claude reviewer memo back: {CLAUDE_BACK}", flush=True)
    else:
        split_marker = "\n## 9. Mechanism Diagnostics"
        if split_marker in back:
            back_results, back_discussion_tail = back.split(split_marker, 1)
            back_discussion = "## 9. Mechanism Diagnostics" + back_discussion_tail
        else:
            midpoint = len(back) // 2
            back_results, back_discussion = back[:midpoint], back[midpoint:]
        review_back_results = get_or_generate(
            CLAUDE_BACK_RESULTS,
            "Claude reviewer memo back-results",
            lambda: claude_review(
                anthropic_key,
                build_claude_segment_review_prompt("results and robustness sections", back_results, addendum),
                max_tokens=2300,
            ),
        )
        review_back_discussion = get_or_generate(
            CLAUDE_BACK_DISCUSSION,
            "Claude reviewer memo mechanism-discussion sections",
            lambda: claude_review(
                anthropic_key,
                build_claude_segment_review_prompt("mechanism, discussion, conclusion sections", back_discussion, addendum),
                max_tokens=2300,
            ),
        )
        review_back = clean_markdown(
            "## Back Results/Robustness Review\n\n"
            + review_back_results
            + "\n\n## Back Mechanism/Discussion Review\n\n"
            + review_back_discussion
        )
        CLAUDE_BACK.write_text(review_back, encoding="utf-8")
    review = clean_markdown(
        "# Claude Reviewer Comments\n\n"
        "## Front-Half Review\n\n"
        + review_front
        + "\n\n## Back-Half Review\n\n"
        + review_back
    )
    CLAUDE_REVIEW.write_text(review, encoding="utf-8")
    log_lines.append(f"- Claude reviewer memo written: `{CLAUDE_REVIEW.relative_to(ROOT)}` ({len(review)} chars).")

    print("Calling DeepSeek for final front-half integration...", flush=True)
    front_final = get_or_generate(
        FINAL_FRONT,
        "DeepSeek final front integration",
        lambda: deepseek_chat(deepseek_key, build_segment_revision_prompt("front", front, ds_draft, review, addendum), max_tokens=4400),
    )
    print("Calling DeepSeek for final back-half integration...", flush=True)
    back_final = get_or_generate(
        FINAL_BACK,
        "DeepSeek final back integration",
        lambda: deepseek_chat(deepseek_key, build_segment_revision_prompt("back", back, ds_draft, review, addendum), max_tokens=5400),
    )
    final = clean_markdown(front_final + "\n\n" + back_final)
    FINAL_MD.write_text(final, encoding="utf-8")
    markdown_to_docx(final, FINAL_DOCX)
    log_lines.append(f"- Final LLM-revised manuscript written: `{FINAL_MD.relative_to(ROOT)}` ({len(final)} chars).")
    log_lines.append(f"- Final DOCX written: `{FINAL_DOCX.relative_to(ROOT)}`.")

    CALL_LOG.write_text("\n".join(log_lines) + "\n", encoding="utf-8")
    print(f"Wrote {FINAL_MD}")
    print(f"Wrote {FINAL_DOCX}")
    print(f"Wrote {CLAUDE_REVIEW}")
    print(f"Wrote {CALL_LOG}")
    return 0


if __name__ == "__main__":
    sys.exit(main())