"""极简 OMML(Word原生公式) 生成器：覆盖本文所需的下标/上标/分式/括号/希腊字母/求和。"""
from lxml import etree

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M = "{%s}" % M_NS
W = "{%s}" % W_NS
NSMAP = {"m": M_NS, "w": W_NS}


def _e(tag, parent=None):
    el = etree.SubElement(parent, M + tag) if parent is not None else etree.Element(M + tag, nsmap=NSMAP)
    return el


def run(text, sty=None):
    """m:r。sty='p' 表示正体（函数名等），默认数学斜体。"""
    r = _e("r")
    if sty:
        rpr = _e("rPr", r)
        s = _e("sty", rpr)
        s.set(M + "val", sty)
    t = _e("t", r)
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return [r]


def T(text):
    return run(text)


def P(text):
    return run(text, sty="p")


def _pack(items):
    """items: list of element-lists -> flat element list"""
    out = []
    for it in items:
        out.extend(it)
    return out


def sub(base, sb):
    """下标 m:sSub"""
    s = _e("sSub")
    e = _e("e", s)
    for x in _pack([base]):
        e.append(x)
    sb_ = _e("sub", s)
    for x in _pack([sb]):
        sb_.append(x)
    return [s]


def sup(base, sp):
    s = _e("sSup")
    e = _e("e", s)
    for x in _pack([base]):
        e.append(x)
    sp_ = _e("sup", s)
    for x in _pack([sp]):
        sp_.append(x)
    return [s]


def frac(num, den):
    f = _e("f")
    n = _e("num", f)
    for x in _pack([num]):
        n.append(x)
    d = _e("den", f)
    for x in _pack([den]):
        d.append(x)
    return [f]


def par(*inner, left="(", right=")"):
    d = _e("d")
    dpr = _e("dPr", d)
    b = _e("begChr", dpr); b.set(M + "val", left)
    en = _e("endChr", dpr); en.set(M + "val", right)
    e = _e("e", d)
    for x in _pack(list(inner)):
        e.append(x)
    return [d]


def hat(base):
    acc = _e("acc")
    accPr = _e("accPr", acc)
    chr_ = _e("chr", accPr); chr_.set(M + "val", "̂")
    e = _e("e", acc)
    for x in _pack([base]):
        e.append(x)
    return [acc]


def bar(base):
    acc = _e("acc")
    accPr = _e("accPr", acc)
    chr_ = _e("chr", accPr); chr_.set(M + "val", "̅")
    e = _e("e", acc)
    for x in _pack([base]):
        e.append(x)
    return [acc]


def seq(*items):
    return _pack(list(items))


def omath(*items):
    om = etree.Element(M + "oMath", nsmap=NSMAP)
    for x in _pack(list(items)):
        om.append(x)
    return om


def add_display_math(doc, items, number=None):
    """居中显示公式段落；number 如 '（1）' 右对齐（用制表位实现）。"""
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    if number:
        tabs = pPr.makeelement(qn("w:tabs"), {})
        for pos, val in [("4500", "center"), ("9000", "right")]:
            t = tabs.makeelement(qn("w:tab"), {})
            t.set(qn("w:val"), val)
            t.set(qn("w:pos"), pos)
            tabs.append(t)
        pPr.append(tabs)
        p.add_run("\t")
    om = omath(*items)
    p._p.append(om)
    if number:
        p.add_run("\t" + number)
    else:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_inline_math(paragraph, items):
    om = omath(*items)
    paragraph._p.append(om)
    return paragraph
