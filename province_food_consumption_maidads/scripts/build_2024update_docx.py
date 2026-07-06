# -*- coding: utf-8 -*-
"""Build From_quantity_to_composition_2024update.docx from the revised source,
preserving the real-per-capita-GDP framing and K=0.137515 US$ story;
update period 2015-2023->2015-2024, obs 279->310, all 11 tables + narrative numbers."""
import json, shutil, os
from docx import Document

ROOT=os.environ.get("PROVINCE_ROOT") or os.path.abspath(os.path.join(os.path.dirname(__file__),"..",".."))
SRC=os.path.join(ROOT,"From_quantity_to_composition_revised.docx")
OUT=os.path.join(ROOT,"ProvinceMAIDADS/From_quantity_to_composition_2024update.docx")
M=json.load(open(os.path.join(ROOT,"ProvinceMAIDADS/scripts/docx_master_all_tables.json")))

shutil.copy(SRC,OUT)
d=Document(OUT)

def set_cell(cell,text):
    p=cell.paragraphs[0]
    if p.runs:
        p.runs[0].text=text
        for r in p.runs[1:]: r.text=""
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs: r.text=""

def upd(ti, rowmap):
    """rowmap: {row_index: {col_index: value}}"""
    t=d.tables[ti]
    for ri,cols in rowmap.items():
        for ci,val in cols.items():
            set_cell(t.rows[ri].cells[ci], val)

import copy as _copy
def ensure_rows(ti, ndata):
    """Ensure table ti has at least ndata data rows (excl. header) by cloning
    the last row's XML (preserves cell formatting)."""
    t=d.tables[ti]
    cur=len(t.rows)-1  # minus header
    last=t.rows[-1]._tr
    for _ in range(ndata-cur):
        new=_copy.deepcopy(last)
        last.addnext(new)
        last=new

# ---------- TABLE 0 (Descriptive) ----------
T0=M['T0']
upd(0,{
 1:{1:T0['budget_m']}, 2:{1:T0['pgdp_usd']}, 3:{1:T0['usd_range']},
 4:{1:T0['exp']}, 5:{1:T0['covered_food_exp']}, 6:{1:T0['residual']},
 7:{1:T0['grain']},8:{1:T0['oil']},9:{1:T0['vegfruit']},10:{1:T0['pork']},
 11:{1:T0['meatother']},12:{1:T0['dairyegg']},
 13:{1:T0['obs'], 2:'31 provinces x 10 years'},
})

# ---------- TABLE 1 (Model fit) ----------
for i,row in enumerate(M['T1']):
    # row=[spec,model,NLL,AIC,BIC,RMSE,OOS]; update cols 2..6 (keep 0,1 labels)
    upd(1,{i+1:{2:row[2],3:row[3],4:row[4],5:row[5],6:row[6]}})

# ---------- TABLE 2 (MAIDADS params) ----------
for i,row in enumerate(M['T2']):
    # row=[label,alpha,delta,tau]; cols1,2,3
    upd(2,{i+1:{1:row[1],2:row[2],3:row[3]}})

# ---------- TABLE 3 (mean-GDP elasticity) ----------
for i,row in enumerate(M['T3']):
    # [group,elas,CI,p,share] -> cols1..4
    upd(3,{i+1:{1:row[1],2:row[2],3:row[3],4:row[4]}})

# ---------- TABLE 4 (US$ node elasticity) ----------
for i,row in enumerate(M['T4']):
    # [group,15k,20k,25k,30k]
    upd(4,{i+1:{1:row[1],2:row[2],3:row[3],4:row[4]}})

# ---------- TABLE 5 (nutrient nodes) ----------
for i,row in enumerate(M['T5']):
    upd(5,{i+1:{1:row[1],2:row[2],3:row[3],4:row[4]}})

# ---------- TABLE 6 (group kg/person/yr) ----------
for i,row in enumerate(M['T6']):
    # [group,2030,2035,2050,natl2050]
    upd(6,{i+1:{1:row[1],2:row[2],3:row[3],4:row[4]}})

# ---------- TABLE 7 (comparison) : update only r1 first cell elasticity list ----------
t7=d.tables[7]
newcmp=("This paper: provinces, 2015-2024, split-pork MAIDADS")
set_cell(t7.rows[1].cells[0], newcmp)
newelas=("Mean-GDP elasticities: staples -0.405 (p=0.018); oils -0.222 (p=0.369); "
 "vegetables and fruits 0.390 (p<0.002); pork 0.317 (p<0.002); non-pork meat/aquatic 0.153 (p=0.004); "
 "dairy and eggs 0.568 (p<0.002).")
set_cell(t7.rows[1].cells[1], newelas)

# ---------- TABLE 8 (own-price) ----------
for i,row in enumerate(M['T8']):
    # [group,Marsh,mp,Hicks,hp]
    upd(8,{i+1:{1:row[1],2:row[2],3:row[3],4:row[4]}})

# ---------- TABLE 9 (feed-grain) ----------
items=M['T9']['items']; tot=M['T9']['totals']
_tci=M['T9']['total_ci_2050']
if isinstance(_tci,str):
    tci=[x.strip() for x in _tci.strip('[]').split(',')]
else:
    tci=[str(x) for x in _tci]
for i,row in enumerate(items):
    lo,hi=[x.strip() for x in row[4].strip('[]').split(',')]
    upd(9,{i+1:{1:row[1],2:row[2],3:row[3],4:lo,5:hi}})
# total row r8
upd(9,{8:{1:tot[0],2:tot[1],3:tot[2],4:tci[0],5:tci[1]}})

# ---------- TABLE 10 (OOS) ----------
# extend to hold the appended 2024-holdout rows, then fill ALL cells (labels+values)
ensure_rows(10, len(M['T10']))
for i,row in enumerate(M['T10']):
    # [spec,model,train,test,rmse,rel,mae] -> fill all 7 columns
    upd(10,{i+1:{0:row[0],1:row[1],2:row[2],3:row[3],4:row[4],5:row[5],6:row[6]}})

# ================= NARRATIVE =================
GLOBAL=[("2015\u20132023","2015\u20132024"),
        ("279 province-year","310 province-year"),
        ("31 provinces x 9 years","31 provinces x 10 years")]

PARA={
 5:[("0.239","0.405"),("0.109","0.222"),("(0.216)","(0.390)"),("(0.251)","(0.317)"),
    ("(0.384)","(0.153)"),("(0.441)","(0.568)"),("approximately 350 million","approximately 393 million")],
 40:[("1,000 province-block draws","450 valid province-block draws")],
 54:[("US$4,200\u2013US$23,900","US$4,000\u2013US$25,400"),
     ("median of approximately US$7,700","median of approximately US$8,000")],
 67:[("-3923","-4248"),("-3999","-4482"),("-7816","-8469"),("-7953","-8919"),
     ("-7762","-8416"),("-7869","-8837"),("0.2913","0.3091"),("0.2813","0.2738"),
     ("0.3006","0.2848"),("0.2805","0.2716")],
 69:[("152.7","466.5"),("0.6265","0.270"),("329.6","205.7"),("0.1546","0.566")],
 75:[("0.801","0.705")],
 81:[("66,734","68,076"),("1,000 successful","450 valid")],
 83:[("66,734","68,076"),("1,000 successful","450 valid")],
 84:[("0.239","0.405"),("0.109","0.222"),("p = 0.004","p = 0.018"),("p = 0.284","p = 0.369")],
 95:[("approximately US$23,900","approximately US$25,400")],
 97:[("1,000 province-block bootstrap parameter draws","450 valid province-block bootstrap parameter draws")],
 98:[("staples (-0.255) and oils (-0.109)","staples (-0.398) and oils (-0.201)"),
     ("p values are 0.058 and 0.284","p values are 0.013 and 0.324"),
     ("vegetables and fruits (0.178), pork (0.191), non-pork meat and aquatic products (0.306), and dairy and eggs (0.329)",
      "vegetables and fruits (0.265), pork (0.184), non-pork meat and aquatic products (0.121), and dairy and eggs (0.346)"),
     ("pork features an elasticity of 0.108, non-pork meat and aquatic products feature an elasticity of 0.211, and dairy and eggs feature an elasticity of 0.203",
      "pork features an elasticity of 0.047, non-pork meat and aquatic products feature an elasticity of 0.079, and dairy and eggs feature an elasticity of 0.140")],
 107:[("393.5 to 413.4 kg","395.4 to 400.4 kg"),
      ("Staples decrease from 121.0 to 93.9","Staples decrease from 113.3 to 84.4"),
      ("vegetables and fruits increase from 172.5 to 198.9","vegetables and fruits increase from 184.3 to 210.1"),
      ("pork increases from 28.8 to 33.2","pork increases from 28.5 to 30.5"),
      ("non-pork meat and aquatic products increase from 32.7 to 41.8","non-pork meat and aquatic products increase from 29.3 to 31.4"),
      ("dairy and eggs increase from 28.5 to 36.7","dairy and eggs increase from 30.6 to 35.6")],
 108:[("vegetables and fruits will account for approximately 277.4 million tons","vegetables and fruits will account for approximately 293.1 million tons"),
      ("staples for 130.9 million tons","staples for 117.7 million tons"),
      ("pork for 46.3 million tons","pork for 42.5 million tons"),
      ("non-pork meat and aquatic products for 58.3 million tons","non-pork meat and aquatic products for 43.8 million tons"),
      ("dairy and eggs for 51.2 million tons","dairy and eggs for 49.7 million tons")],
 124:[("approximately -0.750","approximately -0.676")],
 131:[("2050 point estimate is 349.8 million tons","2050 point estimate is 393.4 million tons"),
      ("accounting for 153.1 million tons, i.e., approximately 44%","accounting for 165.0 million tons, i.e., approximately 42%"),
      ("followed by eggs and poultry at 47.7 million tons each, beef at 39.7 million tons, mutton at 26.4 million tons, aquatic products at 24.2 million tons, and milk at 11.0 million tons",
       "followed by eggs at 62.6 million tons, poultry at 54.1 million tons, beef at 43.8 million tons, mutton at 26.6 million tons, aquatic products at 26.9 million tons, and milk at 14.4 million tons")],
 132:[("132.8 to 253.3 million tons","143.4 to 211.8 million tons")],
 141:[("approximately 349.8 million tons","approximately 393.4 million tons"),
      ("pork contributes approximately 153.1 million tons, accounting for approximately 44%",
       "pork contributes approximately 165.0 million tons, accounting for approximately 42%")],
}

# Full-paragraph rewrites (formatting collapses to first run):
REWRITE={
 68:("The out-of-sample evidence is mixed but strengthens at the most recent horizon. In the main residual-price "
     "specification, the mean food OOS RMSE is essentially unchanged between the two systems in the two earlier "
     "splits: 0.03331 under AIDADS and 0.03337 under MAIDADS in the 2015\u20132020 to 2021\u20132023 split, and 0.03402 "
     "versus 0.03358 in the 2015\u20132022 to 2023 split. In the newest-year holdout, however\u2014fitting 2015\u20132023 and "
     "predicting the 2024 provincial cross-section\u2014MAIDADS clearly outperforms AIDADS, with a mean food OOS RMSE of "
     "0.02931 versus 0.03454, an advantage that holds for every food group. The two systems therefore deliver similar "
     "short-horizon predictions in the pre-2024 splits, while the added flexibility of MAIDADS pays off precisely when "
     "extrapolating to the newest year, so the case for MAIDADS rests on its in-sample likelihood and interpretive "
     "flexibility and, at the most recent horizon, on out-of-sample gains as well. Under the robustness price "
     "specification, the two models are close in all three designs, with MAIDADS marginally behind AIDADS in the 2024 "
     "holdout (0.03503 versus 0.03428)."),
 101:("The nutrient elasticities provide the second key judgment. Covered energy is slightly negative at all four "
      "nodes, with p values well away from significance, thus confirming that aggregate covered calories are effectively "
      "saturated. Plant-source energy is negative and moderately signed (\u22120.279 at US$15,000, p = 0.044), whereas "
      "animal-source energy remains positive and signed across the range (0.199 at US$15,000 declining to 0.070 at "
      "US$30,000, with p values near or below 0.05). Protein and fat elasticities are small and slightly negative and are "
      "imprecisely estimated (p values near or above one-half), reflecting the fact that high-income recomposition is "
      "increasingly a matter of quality, variety, and animal-product mix rather than a large expansion of total modeled "
      "calories or macronutrients. Carbohydrate elasticity is negative (\u22120.278 at US$15,000, p = 0.036), in line "
      "with the decline in staple dependence."),
}

def apply_runs(p, repls):
    """Run-level first (preserves formatting); then paragraph-level for tokens
    that span multiple runs (flattens that paragraph to a single run)."""
    for r in p.runs:
        for a,b in repls:
            if a in r.text:
                r.text=r.text.replace(a,b)
    full="".join(r.text for r in p.runs)
    new=full
    for a,b in repls:
        if a in new:
            new=new.replace(a,b)
    if new!=full:
        if p.runs:
            p.runs[0].text=new
            for r in p.runs[1:]: r.text=""
        else:
            p.add_run(new)

def rewrite_para(p, text):
    if p.runs:
        p.runs[0].text=text
        for r in p.runs[1:]: r.text=""
    else:
        p.add_run(text)

changed=0
for i,p in enumerate(d.paragraphs):
    if i in REWRITE:
        rewrite_para(p, REWRITE[i]); changed+=1
        # still apply global (period) after rewrite via run replace
        apply_runs(p, GLOBAL)
        continue
    repls=list(GLOBAL)+PARA.get(i,[])
    before="".join(r.text for r in p.runs)
    apply_runs(p, repls)
    after="".join(r.text for r in p.runs)
    if before!=after: changed+=1

d.save(OUT)
print("Saved:",OUT)
print("paragraphs changed:",changed)
