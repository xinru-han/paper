# -*- coding: utf-8 -*-
"""Apply the updated results to the manuscript as Word tracked changes.

Reads the original .docx, rewrites the numbers in Tables 2-9 and the affected
sentences in Sections 4-5 as w:ins/w:del (Word "track changes"), saves a
revised .docx. Table 1 is unchanged (the pipeline reproduces it exactly).

Usage: python revise_manuscript.py <original.docx> <revised.docx> <casm_dir>
"""
import os
import sys
import copy

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AUTHOR = "Revision"
DATE = "2026-07-08T00:00:00Z"
_uid = [1000]
warn = []


def _rpr(run_el):
    rpr = run_el.find(qn("w:rPr")) if run_el is not None else None
    return copy.deepcopy(rpr) if rpr is not None else None


def _run(text, rpr, deltext=False):
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement("w:delText" if deltext else "w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _wrap(tag, run_el):
    _uid[0] += 1
    el = OxmlElement(tag)
    el.set(qn("w:id"), str(_uid[0]))
    el.set(qn("w:author"), AUTHOR)
    el.set(qn("w:date"), DATE)
    el.append(run_el)
    return el


def replace_paragraph(p, new):
    """Whole-paragraph tracked replacement; preserves first run formatting."""
    cur = "".join(n.text or "" for n in p._p.iter(qn("w:t")))
    runs = p._p.findall(qn("w:r"))
    rpr = _rpr(runs[0]) if runs else None
    for r in runs:
        p._p.remove(r)
    if cur:
        p._p.append(_wrap("w:del", _run(cur, rpr, deltext=True)))
    if new:
        p._p.append(_wrap("w:ins", _run(new, rpr)))
    return cur


def replace_in_paragraph(p, old, new):
    """Substring tracked replacement inside one run; False if spans runs."""
    for r in p._p.findall(qn("w:r")):
        t = r.find(qn("w:t"))
        if t is None or not t.text or old not in t.text:
            continue
        before, after = t.text.split(old, 1)
        rpr = _rpr(r)
        parent = r.getparent()
        idx = list(parent).index(r)
        parent.remove(r)
        nodes = []
        if before:
            nodes.append(_run(before, rpr))
        nodes.append(_wrap("w:del", _run(old, rpr, deltext=True)))
        if new:
            nodes.append(_wrap("w:ins", _run(new, rpr)))
        if after:
            nodes.append(_run(after, rpr))
        for j, nd in enumerate(nodes):
            parent.insert(idx + j, nd)
        return True
    return False


def edit_body(doc, old, new):
    for p in doc.paragraphs:
        full = "".join(n.text or "" for n in p._p.iter(qn("w:t")))
        if old in full:
            if not replace_in_paragraph(p, old, new):
                replace_paragraph(p, full.replace(old, new))
            return True
    warn.append("正文未匹配: " + old[:30])
    return False


def revise(src, dst, casm_dir):
    doc = Document(src)
    T = doc.tables

    def C(ti, ri, ci, new, pi=0):
        """Set paragraph pi of cell (ri,ci) in table ti to new (tracked)."""
        cell = T[ti].rows[ri].cells[ci]
        p = cell.paragraphs[pi]
        cur = "".join(n.text or "" for n in p._p.iter(qn("w:t")))
        if cur.strip() == new.strip():
            return
        replace_paragraph(p, new)

    # =========================== docx表1 = Table 2 (subgroup) ==============
    C(1, 1, 3, "0.133***"); C(1, 1, 4, "[0.081, 0.185]"); C(1, 1, 5, "95.6")
    C(1, 2, 2, "15"); C(1, 2, 3, "0.089***"); C(1, 2, 4, "[0.033, 0.145]"); C(1, 2, 5, "90.7")
    C(1, 3, 2, "7"); C(1, 3, 3, "0.131***"); C(1, 3, 4, "[0.074, 0.187]"); C(1, 3, 5, "81.7")
    C(1, 5, 3, "0.060**"); C(1, 5, 4, "[0.002, 0.118]")
    C(1, 6, 2, "3"); C(1, 6, 3, "-0.028"); C(1, 6, 4, "[-0.083, 0.027]"); C(1, 6, 5, "93.0")
    C(1, 7, 2, "1"); C(1, 7, 3, "0.176***"); C(1, 7, 4, "[0.137, 0.215]"); C(1, 7, 5, "0.0")
    C(1, 12, 4, "[0.123, 0.290]")

    # =========================== docx表2 = Table 3 (meta-regression) =======
    # cells are two paragraphs: [coef],[ (SE) ]
    reg = {
        1: [("0.562***", "(0.189)"), ("-0.011", "(0.008)"), ("0.252*", "(0.129)")],
        2: [("0.076", "(0.057)"), ("0.187***", "(0.008)"), ("-0.013", "(0.042)")],
        3: [("0.162**", "(0.073)"), ("0.098***", "(0.023)"), ("0.036", "(0.076)")],
        4: [("-0.070***", "(0.025)"), None, ("-0.015", "(0.011)")],
    }
    for ri, cols in reg.items():
        for j, val in enumerate(cols):
            if val is None:
                continue
            coef, se = val
            C(2, ri, j + 1, coef, pi=0)
            C(2, ri, j + 1, se, pi=1)

    # =========================== docx表3 = Table 4 (FAT-PET-PEESE) =========
    C(3, 1, 2, "3.356"); C(3, 1, 3, "0.001")
    C(3, 2, 2, "0.002"); C(3, 2, 3, "0.930")
    C(3, 3, 2, "5.219"); C(3, 3, 3, "0.112")
    C(3, 4, 2, "-0.029"); C(3, 4, 3, "0.176")

    # =========================== docx表4 = Table 5 (robustness) ============
    C(4, 1, 2, "0.133")
    C(4, 2, 2, "0.124"); C(4, 2, 3, "26")
    C(4, 3, 2, "0.143")
    C(4, 4, 2, "0.078")
    C(4, 5, 2, "0.060"); C(4, 6, 2, "0.060"); C(4, 7, 2, "0.061"); C(4, 8, 2, "0.012")
    C(4, 12, 2, "0.078")

    # =========================== docx表5 = Table 6 (parameter mapping) =====
    C(5, 1, 2, "0.089***"); C(5, 1, 3, "0.163（14）")
    C(5, 2, 2, "0.131***"); C(5, 2, 3, "0.164（6）")
    C(5, 3, 3, "0.185（5）")
    C(5, 4, 2, "-0.028")
    C(5, 5, 2, "0.176***")

    # =========================== docx表6 = Table 7 (scenario shifters) =====
    shift = {2: "0.391%", 3: "0.554%", 4: "0.228%",
             6: ("0.458%", "0.066%"), 7: ("0.621%", "0.089%"), 8: ("0.294%", "0.042%"),
             10: ("0.207%", "0.017%"), 11: ("0.449%", "0.036%"), 12: ("0.024%", "0.002%")}
    for ri, v in shift.items():
        if isinstance(v, tuple):
            C(6, ri, 2, v[0]); C(6, ri, 3, v[1])
        else:
            C(6, ri, 2, v)

    # =========================== docx表7 = Table 8 (grain totals) ==========
    g = pd.read_csv(os.path.join(casm_dir, "Table8_grain.csv"),
                    encoding="utf-8-sig")
    order = ["Baseline", "S1-Medium", "S1-High", "S1-Low", "S2-Medium",
             "S2-High", "S2-Low", "S3-Medium", "S3-High", "S3-Low"]
    for i, sc in enumerate(order):
        row = g[g["scenario"] == sc].iloc[0]
        ri = i + 1
        C(7, ri, 1, str(int(row["output_10kt"])))
        if sc != "Baseline":
            C(7, ri, 2, f"{float(row['output_chg_pct']):.2f}")
            C(7, ri, 6, f"{float(row['yield_chg_pct']):.2f}")
        C(7, ri, 3, str(int(row["net_trade_10kt"])))
        C(7, ri, 4, f"{float(row['SSR_pct']):.2f}")
        C(7, ri, 5, f"{float(row['sown_area_100Mmu'])/10:.3f}")

    # =========================== docx表8 = Table 9 (cereal & staple) =======
    s = pd.read_csv(os.path.join(casm_dir, "Table9_cereal_staple.csv"),
                    encoding="utf-8-sig")

    def sgn(x):
        x = int(round(x))
        return f"+{x}" if x > 0 else str(x)

    for i, sc in enumerate(order):
        row = s[s["scenario"] == sc].iloc[0]
        ri = i + 2
        C(8, ri, 1, str(int(row["cereal_output_10kt"])))
        C(8, ri, 2, f"{float(row['cereal_SSR_pct']):.2f}")
        C(8, ri, 3, sgn(row["cereal_net_trade_10kt"]))
        C(8, ri, 4, str(int(row["staple_output_10kt"])))
        C(8, ri, 5, f"{float(row['staple_SSR_pct']):.2f}")
        C(8, ri, 6, sgn(row["staple_net_trade_10kt"]))

    # =========================== body-text edits ==========================
    edit_body(doc,
              "单产维度为95.7%，面积维度为97.1%",
              "单产维度为95.6%，面积维度为97.1%")
    edit_body(doc,
              "在单产维度，农机社会化服务相较于农机资本投入具有显著正向溢价，说明在控制样本规模等方法论特征后，社会化服务路径仍表现出更强的单产提升效应。",
              "在单产维度，综合机械化水平相对农机资本投入具有显著正向溢价，农机社会化服务的溢价为正但不显著，说明在控制样本规模等方法论特征后，区域机械化扩散在单产提升方面表现出更强的相对效应。")
    edit_body(doc,
              "单产维度剔除2条极端样本（文献编号P_22和P_25）后的合并PCC由0.131小幅提高至0.137",
              "单产维度剔除3条极端样本（文献编号P_22、P_25和E_15）后的合并PCC由0.133小幅变动至0.124")
    edit_body(doc,
              "S1-Medium设定每年的单产Shifter=0.434%，S1-High和S1-Low设定每年的单产Shifter分别为0.626%和0.25%",
              "S1-Medium设定每年的单产Shifter=0.391%，S1-High和S1-Low设定每年的单产Shifter分别为0.554%和0.228%")
    edit_body(doc,
              "S2-Medium设定每年的单产Shifter=0.438%，S2-High和S2-Low设定每年的单产Shifter分别为0.606%和0.278%；S2-Medium设定每年的面积Shifter=0.072%，S2-High和S2-Low设定每年的面积Shifter分别为0.098%和0.046%",
              "S2-Medium设定每年的单产Shifter=0.458%，S2-High和S2-Low设定每年的单产Shifter分别为0.621%和0.294%；S2-Medium设定每年的面积Shifter=0.066%，S2-High和S2-Low设定每年的面积Shifter分别为0.089%和0.042%")
    edit_body(doc,
              "S3-Medium设定每年的单产Shifter=0.208%，S3-High和S3-Low设定每年的单产Shifter分别为0.448%和0.024%；S3-Medium设定每年的面积Shifter=0.016%，S3-High和S3-Low设定每年的面积Shifter分别为0.036%和0.002%",
              "S3-Medium设定每年的单产Shifter=0.207%，S3-High和S3-Low设定每年的单产Shifter分别为0.449%和0.024%；S3-Medium设定每年的面积Shifter=0.017%，S3-High和S3-Low设定每年的面积Shifter分别为0.036%和0.002%")
    edit_body(doc,
              "基准情景下，2030年粮食总产量为72530万吨，净进口为14305万吨，自给率为83.53%。在九个农业机械化发展情景下，农机社会化服务路径表现出最高的综合增产潜力：S2-Medium情景下，粮食总产量较基准增加2143万吨，增幅为2.95%；S2-High情景增产2836万吨，增幅达3.91%，为全部情景中最高。",
              "基准情景下，2030年粮食总产量为73377万吨，净进口为12595万吨，自给率为85.35%。在九个农业机械化发展情景下，农机社会化服务路径表现出最高的综合增产潜力：S2-Medium情景下，粮食总产量较基准增加1928万吨，增幅为2.63%；S2-High情景增产2626万吨，增幅达3.58%，为全部情景中最高。")
    edit_body(doc,
              "除S3-Low情景外，其余情景均推动三大谷物自给率超过100%。",
              "除S1-Low、S3-Medium和S3-Low情景外，其余情景均推动三大谷物自给率超过100%。")
    edit_body(doc,
              "在S2-High情景下，谷物自给率提升至102.85%，口粮自给率提升至103.09%，均为全部情景中的最高水平。",
              "在S2-High情景下，谷物自给率提升至101.97%，口粮自给率提升至102.11%，均为全部情景中的最高水平。")

    doc.save(dst)
    print("表格与正文修订完成 ->", dst)
    if warn:
        print("警告:")
        for w in warn:
            print("  " + w)


if __name__ == "__main__":
    revise(sys.argv[1], sys.argv[2], sys.argv[3])
