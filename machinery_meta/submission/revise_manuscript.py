# -*- coding: utf-8 -*-
"""Apply the updated results to the manuscript as Word tracked changes.

Reads the original .docx, rewrites the numbers in Tables 2-9 and the affected
sentences in Sections 4-5 as w:ins/w:del (Word "track changes"), and saves a
revised .docx. Table 1 is unchanged (the updated pipeline reproduces it
exactly).

Usage: python revise_manuscript.py <original.docx> <revised.docx>
"""
import sys
import copy
import datetime

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AUTHOR = "Revision (updated results)"
DATE = "2026-07-08T00:00:00Z"
_uid = [1000]


def _rpr_from(run_el):
    rpr = run_el.find(qn("w:rPr")) if run_el is not None else None
    return copy.deepcopy(rpr) if rpr is not None else None


def _make_run(text, rpr, deltext=False):
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement("w:delText" if deltext else "w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _wrap(tag, run_el):
    _uid[0] += 1
    el = OxmlElement(tag)
    el.set(qn("w:id"), str(_uid[0]))
    el.set(qn("w:author"), AUTHOR)
    el.set(qn("w:date"), DATE)
    el.append(run_el)
    return el


def replace_paragraph(p, old, new):
    """Whole-paragraph tracked replacement (used for table cells & captions).

    Deletes every existing run (as w:del) and inserts `new` (as w:ins).
    Returns True if `old` matched the current paragraph text."""
    cur = "".join(node.text or "" for node in p._p.iter(qn("w:t")))
    if old is not None and old != cur:
        return False
    runs = p._p.findall(qn("w:r"))
    rpr = _rpr_from(runs[0]) if runs else None
    for r in runs:
        p._p.remove(r)
    # also drop any pre-existing empty w:ins/w:del leftovers
    if cur:
        p._p.append(_wrap("w:del", _make_run(cur, rpr, deltext=True)))
    if new:
        p._p.append(_wrap("w:ins", _make_run(new, rpr)))
    return True


def set_cell(cell, new, old=None):
    """Tracked replacement of a table cell's text."""
    p = cell.paragraphs[0]
    return replace_paragraph(p, old, new)


def replace_in_paragraph(p, old, new):
    """Substring tracked replacement inside a body paragraph.

    Works when `old` lies within a single run (true for the numeric edits);
    returns False otherwise so the caller can fall back / warn."""
    for r in p._p.findall(qn("w:r")):
        t = r.find(qn("w:t"))
        if t is None or not t.text or old not in t.text:
            continue
        before, after = t.text.split(old, 1)
        rpr = _rpr_from(r)
        parent = r.getparent()
        idx = list(parent).index(r)
        parent.remove(r)
        new_nodes = []
        if before:
            new_nodes.append(_make_run(before, rpr))
        new_nodes.append(_wrap("w:del", _make_run(old, rpr, deltext=True)))
        if new:
            new_nodes.append(_wrap("w:ins", _make_run(new, rpr)))
        if after:
            new_nodes.append(_make_run(after, rpr))
        for j, node in enumerate(new_nodes):
            parent.insert(idx + j, node)
        return True
    return False


def edit_body(doc, old, new, limit=1):
    """Replace `old`->`new` in body paragraphs (tracked). Returns hit count."""
    n = 0
    for p in doc.paragraphs:
        if n >= limit:
            break
        full = "".join(node.text or "" for node in p._p.iter(qn("w:t")))
        if old in full:
            if replace_in_paragraph(p, old, new):
                n += 1
            else:
                # fall back: whole-paragraph tracked rewrite
                replace_paragraph(p, full, full.replace(old, new))
                n += 1
    return n


# ============================================================ edit spec
def revise(src, dst):
    doc = Document(src)
    T = doc.tables
    warn = []

    def cell(ti, ri, ci, new):
        c = T[ti].rows[ri].cells[ci]
        old = c.text.strip()
        if old == new:
            return
        if not set_cell(c, new, old):
            warn.append(f"表{ti}[{ri},{ci}] 期望'{old}'未匹配")

    # ---- docx表1 = manuscript Table 2 (分路径) ----
    # 列: 维度|路径|k|合并PCC|CI|I2   行1..12
    # 单产全样本
    cell(1, 1, 2, "29"); cell(1, 1, 3, "0.133***"); cell(1, 1, 4, "[0.081, 0.185]"); cell(1, 1, 5, "95.6")
    # 单产MCI
    cell(1, 2, 2, "15"); cell(1, 2, 3, "0.089***"); cell(1, 2, 4, "[0.033, 0.145]"); cell(1, 2, 5, "90.7")
    # 单产AMS
    cell(1, 3, 2, "7"); cell(1, 3, 3, "0.131***"); cell(1, 3, 4, "[0.074, 0.187]"); cell(1, 3, 5, "81.7")
    # 单产AML 不变(7,0.221**,[0.020,0.422],98.5)
    # 面积全样本
    cell(1, 5, 3, "0.060**"); cell(1, 5, 4, "[0.002, 0.118]")
    # 面积MCI
    cell(1, 6, 2, "3"); cell(1, 6, 3, "-0.028"); cell(1, 6, 4, "[-0.083, 0.027]"); cell(1, 6, 5, "93.0")
    # 面积AMS
    cell(1, 7, 2, "1"); cell(1, 7, 3, "0.176***"); cell(1, 7, 4, "[0.137, 0.215]"); cell(1, 7, 5, "0.0")
    # 面积AML 不变(3,0.112***,[0.050,0.174],86.0)
    # 效率行 全不变；效率AML CI 0.124->0.123 微调
    cell(1, 12, 4, "[0.123, 0.290]")

    # ---- docx表2 = Table 3 (元回归) 列: 变量|单产|面积|效率 ----
    cell(2, 1, 1, "0.562*** (0.189)"); cell(2, 1, 2, "-0.011 (0.008)"); cell(2, 1, 3, "0.252* (0.129)")   # 常数项
    cell(2, 2, 1, "0.076 (0.057)"); cell(2, 2, 2, "0.187*** (0.008)"); cell(2, 2, 3, "-0.013 (0.042)")     # 农机社会化服务
    cell(2, 3, 1, "0.162** (0.073)"); cell(2, 3, 2, "0.098*** (0.023)"); cell(2, 3, 3, "0.036 (0.076)")    # 综合机械化水平
    cell(2, 4, 1, "-0.070*** (0.025)"); cell(2, 4, 3, "-0.015 (0.011)")                                     # LogN

    # ---- docx表3 = Table 4 (FAT-PET-PEESE) 列: 维度|检验|系数|值 ----
    cell(3, 1, 2, "3.356"); cell(3, 1, 3, "0.001")      # 单产FAT
    cell(3, 2, 2, "0.002"); cell(3, 2, 3, "0.930")      # 单产PET
    cell(3, 3, 2, "5.219"); cell(3, 3, 3, "0.112")      # 面积FAT
    cell(3, 4, 2, "-0.029"); cell(3, 4, 3, "0.176")     # 面积PET
    # 效率FAT/PET/PEESE 不变(2.683/0.062/0.072)

    # ---- docx表4 = Table 5 (稳健性) 列: 维度|权重设定|合并PCC|样本量 ----
    cell(4, 1, 2, "0.133")                               # 单产 随机效应
    cell(4, 2, 2, "0.124"); cell(4, 2, 3, "26")          # 单产 IQR
    cell(4, 3, 2, "0.143")                               # 单产 简单平均
    cell(4, 4, 2, "0.078")                               # 单产 样本量加权
    cell(4, 5, 2, "0.060")                               # 面积 随机效应
    cell(4, 6, 2, "0.060")                               # 面积 IQR
    cell(4, 7, 2, "0.061")                               # 面积 简单平均
    cell(4, 8, 2, "0.012")                               # 面积 样本量加权
    # 效率 随机/IQR/简单 不变(0.172/0.172/0.181)
    cell(4, 12, 2, "0.078")                              # 效率 样本量加权

    doc.save(dst)
    if warn:
        print("WARN:")
        for w in warn:
            print("  " + w)
    else:
        print("表格改写全部匹配成功")
    return dst


if __name__ == "__main__":
    revise(sys.argv[1], sys.argv[2])
