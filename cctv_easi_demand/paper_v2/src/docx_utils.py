"""Shared helpers for building the Food Policy manuscript with python-docx.

Equations go through latex2mathml -> mathml2omml -> native Word OMML.
Known upstream bug: \bar{} produces an unclosed <m:groupChrPr>; we both avoid
\bar (use \overline) and patch the output defensively.
"""
import re
import latex2mathml.converter
import mathml2omml
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

M_NS = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'

_GROUPCHR_FIX = re.compile(r'(<m:groupChrPr>(?:(?!</m:groupChrPr>).)*?)</m:groupChr>(<m:e>)')


def latex_to_omml(tex: str) -> str:
    mml = latex2mathml.converter.convert(tex)
    omml = mathml2omml.convert(mml)
    omml = _GROUPCHR_FIX.sub(r'\1</m:groupChrPr>\2', omml)
    return omml.replace('<m:oMath>', f'<m:oMath {M_NS}>', 1)


def add_equation(doc, tex: str, number: str | None = None):
    """Display equation, centered; optional right-aligned (number) via tab."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p._p.append(parse_xml(latex_to_omml(tex)))
    if number:
        run = p.add_run(f"  ({number})")
        run.font.size = Pt(11)
    return p


def add_inline_math(paragraph, tex: str):
    paragraph._p.append(parse_xml(latex_to_omml(tex)))
    return paragraph


def set_base_styles(doc, font="Times New Roman", size=11, line_spacing=1.5):
    st = doc.styles["Normal"]
    st.font.name = font
    st.font.size = Pt(size)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    st.paragraph_format.line_spacing = line_spacing
    st.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_table(doc, header, rows, col_widths=None, note=None, font_size=9):
    """Simple bordered journal-style table. header: list[str]; rows: list[list]."""
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Table Grid"
    for j, htxt in enumerate(header):
        cell = t.cell(0, j)
        cell.text = str(htxt)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(font_size)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i + 1, j)
            cell.text = "" if val is None else str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    if note:
        p = doc.add_paragraph()
        run = p.add_run(note)
        run.font.size = Pt(8)
        run.italic = True
    return t


def fmt(x, d=3):
    try:
        return f"{float(x):.{d}f}"
    except (TypeError, ValueError):
        return str(x)


def stars(est, se):
    try:
        z = abs(float(est) / float(se))
    except (TypeError, ValueError, ZeroDivisionError):
        return ""
    if z >= 2.576:
        return "***"
    if z >= 1.960:
        return "**"
    if z >= 1.645:
        return "*"
    return ""
